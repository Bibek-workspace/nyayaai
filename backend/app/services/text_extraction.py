"""Extract plain text from uploaded documents.

Strategy:
  1. PDF — try pypdf first (fast, works for digital PDFs).
  2. If text is empty or below a threshold, fall back to OCR (Tesseract via pdf2image).
  3. DOCX — python-docx.
  4. Plain text / unknown — decode as UTF-8 with replace fallback.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

import pytesseract
from pdf2image import convert_from_bytes
from pypdf import PdfReader

from app.core.config import settings

logger = logging.getLogger(__name__)
pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

MIN_USEFUL_TEXT_CHARS = 200


class TextExtractionResult:
    def __init__(self, text: str, ocr_used: bool):
        self.text = text
        self.ocr_used = ocr_used


def _extract_pdf_native(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:
                logger.warning("Page extract failed: %s", e)
        return "\n".join(parts).strip()
    except Exception as e:
        logger.warning("Native PDF parse failed: %s", e)
        return ""


def _extract_pdf_ocr(data: bytes) -> str:
    try:
        pages = convert_from_bytes(data, dpi=200)
    except Exception as e:
        logger.error("pdf2image failed: %s", e)
        return ""
    chunks = []
    for img in pages:
        try:
            chunks.append(pytesseract.image_to_string(img, lang=settings.OCR_LANGUAGES))
        except Exception as e:
            logger.warning("Tesseract page failed: %s", e)
    return "\n".join(chunks).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    tables = []
    for table in doc.tables:
        for row in table.rows:
            tables.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + tables).strip()


def extract_text(data: bytes, filename: str, mime_type: str) -> TextExtractionResult:
    name_lower = filename.lower()
    suffix = Path(filename).suffix.lower()
    ocr_used = False

    if mime_type == "application/pdf" or suffix == ".pdf":
        text = _extract_pdf_native(data)
        if len(text) < MIN_USEFUL_TEXT_CHARS:
            logger.info("Native PDF text too short (%d chars); falling back to OCR", len(text))
            ocr_text = _extract_pdf_ocr(data)
            if len(ocr_text) > len(text):
                text = ocr_text
                ocr_used = True
        return TextExtractionResult(text=text, ocr_used=ocr_used)

    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return TextExtractionResult(text=_extract_docx(data), ocr_used=False)

    if mime_type.startswith("image/"):
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        return TextExtractionResult(
            text=pytesseract.image_to_string(img, lang=settings.OCR_LANGUAGES).strip(),
            ocr_used=True,
        )

    # Plain text fallback
    try:
        return TextExtractionResult(text=data.decode("utf-8", errors="replace").strip(), ocr_used=False)
    except Exception:
        return TextExtractionResult(text="", ocr_used=False)
